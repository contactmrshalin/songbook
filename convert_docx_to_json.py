#!/usr/bin/env python3
"""
DOCX → songs.json extractor (tuned for your songbook template)

Features:
- Extract title + info (supports your header table layout)
- Extract notation triplets:
    Lyrics line (bold-ish / texty) →
    Indian line

Note:
- We intentionally do NOT store `western` or `tokens` in songs.json anymore.
- The build pipeline derives Western display and MusicXML tokens from the Indian line at render time.

Outputs each song entry as:
{
  "id", "title", "info", "thumbnail", "background",
  "sections":[{"name":"STHAYI","lines":[{lyrics, indian}]}]
}

You can later edit thumbnail/background paths in songs.json.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document


# -----------------------------
# Helpers
# -----------------------------
def slugify(s: str) -> str:
    s = s.lower().strip()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    return s.strip("-") or "song"


def normalize_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def dedupe_keep_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in items:
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out


# -----------------------------
# Notation detection
# -----------------------------
def is_probably_western_line(text: str) -> bool:
    # Western lines often contain A-G, #/♭, octave marks, etc.
    t = text.strip()
    if not t:
        return False
    if re.search(r"[A-G][#♭]?\d?", t):
        return True
    if re.search(r"\b[A-G]\b", t):
        return True
    return False


def is_probably_sargam_line(text: str) -> bool:
    # Indian sargam tokens + octave marks; allow lowercase komal
    t = text.strip()
    if not t:
        return False
    if re.search(r"\b[SRGmMPDNrgdn]\b", t):
        return True
    if re.search(r"[SRGmMPDNrgdn][’']", t):
        return True
    if "~" in t or "(k" in t.lower() or "(K" in t:
        return True
    return False


def is_probably_lyrics_line(text: str) -> bool:
    # Lyrics: more alphabetic content, fewer note tokens
    t = text.strip()
    if not t:
        return False
    letters = len(re.findall(r"[A-Za-zअ-ह]", t))
    notes = len(re.findall(r"\b[SRGmMPDNrgdn]\b", t))
    return letters >= 6 and notes <= 2


# -----------------------------
# Tokenizer (AUTO-DETECT ornaments)
# -----------------------------
KOMAL_INLINE_RE = re.compile(r"^([RGDN])\((k|K)\)$")       # D(k) -> d
NOTE_CORE_RE = re.compile(r"^[SRGmMPDNrgdn]$")            # single
MEEND_RE = re.compile(r"^([SRGmMPDNrgdn])~([SRGmMPDNrgdn])$")
KAN_RE = re.compile(r"^\(([SRGmMPDNrgdn])\)([SRGmMPDNrgdn])([’']|\.|)?$")
OCT_RE = re.compile(r"^([,]?[SRGmMPDNrgdn])([’']|\.)?$")  # supports ,S and S' or S.

def _norm_note(token: str) -> str:
    """
    Normalize komal inline D(k)->d, keep apostrophe as ' and dot as .
    Keep low-octave comma if it is the first char.
    """
    t = token.strip()
    t = t.replace("’", "'")

    low = t.startswith(",")
    if low:
        t = t[1:]

    octv = ""
    if t.endswith("'"):
        octv = "'"
        t = t[:-1]
    elif t.endswith("."):
        octv = "."
        t = t[:-1]

    m = KOMAL_INLINE_RE.match(t)
    if m:
        t = m.group(1).lower()  # komal
    # if already lowercase komal, keep

    if low:
        return "," + t + octv
    return t + octv


def tokenize_indian(indian_line: str) -> List[str]:
    """
    Convert messy Indian line text into token list:
      - note tokens: S R G m M P D N (+ low ',' + octave ' or .)
      - ornaments:
          meend: G~R
          kan: (R)G
          hold: suffix ":" (canonical) created from …, ..., ---, or ':'.

    This is robust to your typical separators: .., …, |, commas in text.
    """
    s = indian_line.strip()
    if not s:
        return []

    # normalize separators
    s = s.replace("|", " ")
    s = s.replace("…", " ... ")
    s = re.sub(r"\s+", " ", s)

    # break by spaces, but keep ornament chunks
    parts = s.split(" ")
    out: List[str] = []

    for p in parts:
        p = p.strip()
        if not p:
            continue

        # strip common dot separators used like R..G.. etc
        p = p.replace("..", "")
        p = p.replace(".", "") if p.count(".") >= 2 else p  # heavy dots become separators

        # detect hold patterns -> canonical ":"
        hold = ""
        if "..." in p or re.search(r"--+$", p) or re.search(r"\.{3,}$", p) or p.endswith(":"):
            hold = ":"
            p = re.sub(r":$", "", p)
            p = re.sub(r"--+$", "", p)
            p = re.sub(r"\.{3,}$", "", p)
            p = p.replace("...", "")
            p = p.strip()

        # KAN: (R)G or (R)G'
        km = KAN_RE.match(p)
        if km:
            kan = _norm_note(km.group(1))
            main = _norm_note(km.group(2) + (km.group(3) or ""))
            out.append(f"({kan}){main}")
            if hold:
                out.append(main + ":")
            continue

        # MEEND: G~R
        mm = MEEND_RE.match(p)
        if mm:
            out.append(f"{_norm_note(mm.group(1))}~{_norm_note(mm.group(2))}")
            continue

        # Komal inline: D(k)
        if KOMAL_INLINE_RE.match(p):
            out.append(_norm_note(p) + hold)
            continue

        # Simple note with octave/low
        om = OCT_RE.match(p.replace("’", "'"))
        if om:
            out.append(_norm_note(p) + hold)
            continue

        # Some docs contain bracketed patterns like RG{MGRS}S
        # We expand {MGRS} into M G R S
        if "{" in p and "}" in p:
            pre = p.split("{", 1)[0]
            mid = p.split("{", 1)[1].split("}", 1)[0]
            post = p.split("}", 1)[1] if "}" in p else ""
            for ch in (pre + mid + post):
                if ch in "SRGmMPDNrgdn":
                    out.append(_norm_note(ch))
            continue

        # Otherwise ignore unknown fragments
        # (keeps extractor robust against random words)
        continue

    return out


# -----------------------------
# Extract title + info
# -----------------------------
def extract_title_and_info(doc: Document) -> Tuple[str, List[str]]:
    title: Optional[str] = None
    info: List[str] = []

    # Prefer first table (your template)
    if doc.tables:
        t = doc.tables[0]
        collected: List[str] = []
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    s = p.text.strip()
                    if s:
                        collected.append(s)
        collected = dedupe_keep_order(collected)
        if collected:
            title = collected[0]
            for x in collected[1:]:
                if len(x) < 140:
                    info.append(x)

    # Fallback: first non-empty paragraph as title
    if not title:
        for p in doc.paragraphs[:12]:
            s = p.text.strip()
            if s:
                title = s
                break

    if not title:
        title = "Untitled Song"

    # Also capture some early paragraphs as info (avoid duplicates)
    for p in doc.paragraphs[:40]:
        s = p.text.strip()
        if not s or s == title:
            continue
        if len(s) > 140:
            continue
        if "format:" in s.lower():
            continue
        info.append(s)

    info = dedupe_keep_order(info)[:10]
    return title.strip(), info


# -----------------------------
# Extract triplets (Lyrics / Indian)
# -----------------------------
def extract_triplets(doc: Document) -> List[Dict[str, Any]]:
    paras = [normalize_spaces(p.text) for p in doc.paragraphs if p.text and p.text.strip()]
    out: List[Dict[str, Any]] = []

    i = 0
    while i < len(paras):
        a = paras[i]
        b = paras[i + 1] if i + 1 < len(paras) else ""
        c = paras[i + 2] if i + 2 < len(paras) else ""

        # Common pattern: lyrics then Indian then Western
        if is_probably_lyrics_line(a) and is_probably_sargam_line(b):
            out.append({
                "lyrics": a,
                "indian": b,
            })
            i += 3
            continue

        # Pattern: sometimes lyrics line is empty and they just list notations; still capture
        if is_probably_sargam_line(a) and is_probably_western_line(b):
            out.append({
                "lyrics": "",
                "indian": a,
            })
            i += 2
            continue

        i += 1

    return out


def extract_song_from_docx(docx_path: Path) -> Dict[str, Any]:
    doc = Document(str(docx_path))
    title, info = extract_title_and_info(doc)
    lines = extract_triplets(doc)

    song_id = slugify(title)

    # If the doc itself contains explicit headings, user can edit later;
    # we default to single STHAYI section to keep it simple.
    sections = []
    if lines:
        sections = [{"name": "STHAYI", "lines": lines}]

    return {
        "id": song_id,
        "title": title,
        "export": True,
        "info": info,
        "thumbnail": "",
        "background": "",
        "sections": sections,
    }


def save_song_to_dir(root: Path, song: Dict[str, Any]) -> Path:
    """Save a song to ``songs/<id>.json`` and update ``book.json`` song_order.

    Delegates to the shared ``load_songs`` module when available, with an
    inline fallback so this script can still run standalone.
    """
    try:
        from load_songs import save_song, load_song_order, save_book_meta, load_book_meta
        path = save_song(root, song)
        # Ensure song appears in the ordering list.
        order = load_song_order(root)
        song_id = str(song["id"])
        if song_id not in order:
            order.append(song_id)
            title, meta = load_book_meta(root)
            save_book_meta(root, title, meta, order)
        return path
    except ImportError:
        pass

    # Inline fallback (no load_songs module available).
    songs_dir = root / "songs"
    songs_dir.mkdir(exist_ok=True)

    song_id = str(song["id"])
    song_path = songs_dir / f"{song_id}.json"
    song_path.write_text(json.dumps(song, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    book_path = root / "book.json"
    if book_path.exists():
        book = json.loads(book_path.read_text(encoding="utf-8"))
    else:
        book = {"book_title": "My Songbook", "book_meta": {}, "song_order": []}

    order = book.get("song_order", [])
    if song_id not in order:
        order.append(song_id)
        book["song_order"] = order
        book_path.write_text(json.dumps(book, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    return song_path


def merge_songs_into_json(json_path: Path, new_songs: List[Dict[str, Any]]) -> None:
    if json_path.exists():
        data = json.loads(json_path.read_text(encoding="utf-8"))
    else:
        data = {"book_title": "My Songbook", "songs": []}

    songs = data.get("songs", [])
    if not isinstance(songs, list):
        songs = []

    by_id = {s["id"]: s for s in songs if isinstance(s, dict) and "id" in s}

    existing_order = [s["id"] for s in songs if isinstance(s, dict) and "id" in s]

    for s in new_songs:
        # Default: include songs in exports unless user disables later.
        if isinstance(s, dict) and "export" not in s:
            s["export"] = True
        by_id[s["id"]] = s
        if s["id"] not in existing_order:
            existing_order.append(s["id"])

    data["songs"] = [by_id[i] for i in existing_order if i in by_id]

    # Ensure all existing songs have export flag when writing.
    for s in data["songs"]:
        if isinstance(s, dict) and "export" not in s:
            s["export"] = True

    json_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


# -----------------------------
# CLI
# -----------------------------
def main() -> None:
    ap = argparse.ArgumentParser(description="Extract songs from DOCX into songs.json entries (auto ornaments)")
    ap.add_argument("docx", nargs="+", help="DOCX file(s)")
    ap.add_argument("--out", default="songs.json", help="songs.json output path")
    args = ap.parse_args()

    new_songs = [extract_song_from_docx(Path(p)) for p in args.docx]

    out_path = Path(args.out)
    root = out_path.parent
    songs_dir = root / "songs"

    # Use per-song layout if it exists
    if (songs_dir.is_dir() and any(songs_dir.glob("*.json"))) or (root / "book.json").exists():
        for s in new_songs:
            save_song_to_dir(root, s)
        print(f"Saved {len(new_songs)} song(s) to {songs_dir}/")
    else:
        merge_songs_into_json(out_path, new_songs)
        print(f"Added/updated {len(new_songs)} song(s) into {args.out}")


if __name__ == "__main__":
    main()
