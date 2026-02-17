#!/usr/bin/env python3
"""
Songbook Pipeline
Generates:
- Mobile-friendly PDF (single-column, clickable index)
- EPUB (reflowable, clickable TOC, optional cover + metadata)
- MusicXML per song (MuseScore-friendly) with:
  - Indian sargam as lyric line 1
  - Western pitch+octave as lyric line 2
  - Ornaments auto-rendered from tokens: meend (slur), kan (grace), hold (longer duration)

Source of truth: songs.json
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import zipfile
import math
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# -----------------------------
# Notation mapping (centralized)
# -----------------------------
_ROOT = Path(__file__).resolve().parent
_NOTATION_MAPPING_PATH = _ROOT / "notation_mapping.json"


def _default_notation_mapping() -> Dict[str, Any]:
    """
    Fallback mapping used if notation_mapping.json is missing/unreadable.
    Mirrors the repo defaults: SA=C, komal=flats on Re/Ga/Dha/Ni, tivra Ma = F#.
    """
    return {
        "octave_markers": {"low": ".", "middle": "", "high": "'"},
        "accidental_markers": {"komal": "(k)", "tivra": "(T)"},
        "word_to_token": {"Sa": "S", "Re": "R", "Ga": "G", "Ma": "m", "Pa": "P", "Dha": "D", "Ni": "N"},
        "komal_word_to_token": {"Re": "r", "Ga": "g", "Dha": "d", "Ni": "n"},
        "tivra_word_to_token": {"Ma": "M"},
        "token_to_western": {
            "S": {"step": "C", "alter": 0},
            "R": {"step": "D", "alter": 0},
            "G": {"step": "E", "alter": 0},
            "m": {"step": "F", "alter": 0},
            "M": {"step": "F", "alter": 1},
            "P": {"step": "G", "alter": 0},
            "D": {"step": "A", "alter": 0},
            "N": {"step": "B", "alter": 0},
            "r": {"step": "D", "alter": -1},
            "g": {"step": "E", "alter": -1},
            "d": {"step": "A", "alter": -1},
            "n": {"step": "B", "alter": -1},
        },
    }


def _load_notation_mapping() -> Dict[str, Any]:
    try:
        if _NOTATION_MAPPING_PATH.exists():
            return json.loads(_NOTATION_MAPPING_PATH.read_text(encoding="utf-8"))
    except Exception:
        pass
    return _default_notation_mapping()


_NOTATION = _load_notation_mapping()
_TOKEN_TO_WESTERN: Dict[str, Tuple[str, int]] = {
    k: (str(v.get("step", "")), int(v.get("alter", 0)))
    for k, v in (_NOTATION.get("token_to_western", {}) or {}).items()
}

# DOCX deps (optional until DOCX generation is invoked)
try:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore
except ModuleNotFoundError:
    Document = None  # type: ignore
    Inches = None  # type: ignore
    Pt = None  # type: ignore

# PDF deps (optional until PDF generation is invoked)
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A5, letter
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ModuleNotFoundError as e:
    # Allow non-PDF functionality (e.g., MusicXML) to still be usable.
    # We'll surface a clear error only when PDF generation is requested.
    missing = getattr(e, "name", "") or str(e)
    if "reportlab" in missing:
        canvas = None  # type: ignore
        A5 = letter = None  # type: ignore
        inch = None  # type: ignore
        ImageReader = None  # type: ignore
        pdfmetrics = None  # type: ignore
        TTFont = None  # type: ignore
    else:
        raise


def _require_reportlab() -> None:
    if canvas is not None:
        return
    raise SystemExit(
        "Missing dependency: reportlab\n\n"
        "To enable PDF generation, install dependencies (recommended in a venv):\n"
        "  python3 -m venv .venv\n"
        "  source .venv/bin/activate\n"
        "  pip install reportlab pillow python-docx tkinterdnd2\n\n"
        "Or run:\n"
        "  ./setup.sh\n"
    )


def _require_docx() -> None:
    if Document is not None:
        return
    raise SystemExit(
        "Missing dependency: python-docx\n\n"
        "To enable Word (DOCX) generation, install dependencies:\n"
        "  python3 -m venv .venv\n"
        "  source .venv/bin/activate\n"
        "  pip install python-docx\n\n"
        "Or run:\n"
        "  ./setup.sh\n"
    )


# -----------------------------
# Utilities
# -----------------------------
def xesc(t: str) -> str:
    return (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;").replace("'", "&apos;"))


def _boolish(v: Any, *, default: bool = True) -> bool:
    """
    Convert a loose JSON value into a bool.
    Accepts bool/int/str like: true/false, 1/0, "yes"/"no".
    """
    if v is None:
        return default
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return bool(v)
    if isinstance(v, str):
        s = v.strip().lower()
        if s in ("true", "t", "1", "yes", "y", "on"):
            return True
        if s in ("false", "f", "0", "no", "n", "off"):
            return False
    return default


def slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    s = s.strip("-")
    return s or "song"


def safe_path(base_dir: Path, rel: str) -> Optional[Path]:
    """
    Resolve a relative path safely inside base_dir. Returns None if missing/outside.
    """
    if not rel:
        return None
    p = (base_dir / rel).resolve()
    try:
        p.relative_to(base_dir.resolve())
    except ValueError:
        return None
    return p if p.exists() else None


def normalize_export_flags(json_path: Path) -> int:
    """
    Ensure every song object has an explicit boolean `export` key.
    Does NOT overwrite existing values. Returns count of songs updated.
    """
    data = json.loads(json_path.read_text(encoding="utf-8"))
    songs = data.get("songs", [])
    if not isinstance(songs, list):
        return 0

    changed = 0
    for s in songs:
        if not isinstance(s, dict):
            continue
        if "export" not in s:
            s["export"] = True
            changed += 1
            continue

        # normalize to strict bool if user used strings/ints
        b = _boolish(s.get("export", True), default=True)
        if s.get("export") is not b:
            s["export"] = b
            changed += 1

    if changed:
        json_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return changed


@dataclass(frozen=True)
class PdfFontPack:
    regular: str
    bold: str
    italic: str
    bold_italic: str
    unicode_ok: bool


def _register_unicode_font_pack(base_dir: Path) -> PdfFontPack:
    """
    Register a Unicode-capable font pack (regular/bold/italic/bold-italic) for PDF rendering.
    This fixes missing glyphs in ReportLab's built-in fonts (e.g., Latin Extended like Ṡ/Ṙ).

    Strategy:
    - Prefer project-local fonts in ./fonts (DejaVuSans*.ttf shipped with this repo).
    - Fall back to common system font locations.
    - If none available, fall back to built-in Times fonts (limited Unicode).
    """
    # Default (safe) fallback: built-in fonts (NOT fully Unicode).
    fallback = PdfFontPack(
        regular="Times-Roman",
        bold="Times-Bold",
        italic="Times-Italic",
        bold_italic="Times-BoldItalic",
        unicode_ok=False,
    )

    fonts_dir = (base_dir / "fonts")

    # Prefer repo-shipped DejaVu Sans family (good Unicode coverage, portable).
    def p(rel: str) -> Path:
        return (fonts_dir / rel).resolve()

    preferred: List[Tuple[Path, Path, Path, Path]] = []
    if fonts_dir.exists():
        preferred.append((
            p("DejaVuSans.ttf"),
            p("DejaVuSans-Bold.ttf"),
            p("DejaVuSans-Oblique.ttf"),
            p("DejaVuSans-BoldOblique.ttf"),
        ))

    # Common system locations (best-effort). We may not have the full family.
    system_candidates: List[Tuple[Path, Path, Path, Path]] = [
        # Linux (Debian/Ubuntu)
        (
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
        ),
        # Linux (various)
        (
            Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
            Path("/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf"),
            Path("/usr/share/fonts/dejavu/DejaVuSans-BoldOblique.ttf"),
        ),
        # macOS (Supplemental)
        (
            Path("/System/Library/Fonts/Supplemental/DejaVu Sans.ttf"),
            Path("/System/Library/Fonts/Supplemental/DejaVu Sans Bold.ttf"),
            Path("/System/Library/Fonts/Supplemental/DejaVu Sans Oblique.ttf"),
            Path("/System/Library/Fonts/Supplemental/DejaVu Sans Bold Oblique.ttf"),
        ),
    ]

    for reg_fp, bold_fp, ital_fp, bold_ital_fp in (preferred + system_candidates):
        try:
            if not (reg_fp.exists() and reg_fp.is_file()):
                continue

            def reg_font(fp: Path, suffix: str) -> Optional[str]:
                if not (fp.exists() and fp.is_file()):
                    return None
                name = f"SongbookText-{suffix}-{fp.stem}"
                if name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(name, str(fp)))
                return name

            reg_name = reg_font(reg_fp, "Regular")
            if not reg_name:
                continue

            bold_name = reg_font(bold_fp, "Bold") or reg_name
            ital_name = reg_font(ital_fp, "Italic") or reg_name
            bold_ital_name = reg_font(bold_ital_fp, "BoldItalic") or bold_name

            return PdfFontPack(
                regular=reg_name,
                bold=bold_name,
                italic=ital_name,
                bold_italic=bold_ital_name,
                unicode_ok=True,
            )
        except Exception:
            continue

    return fallback


def _western_pdf_safe(s: str, unicode_font_available: bool) -> str:
    """
    If we can't load a Unicode font, replace musical accidentals with ASCII
    so the PDF remains readable.
    """
    if unicode_font_available:
        return s
    return (
        s.replace("♭", "b")
         .replace("♯", "#")
    )


def _pdf_safe_text(s: str, *, emoji_font_available: bool) -> str:
    """
    ReportLab + typical TTF fonts (e.g., DejaVu Sans) generally do NOT render
    color emoji glyphs like 🎤 ✍️ 🎼. To avoid missing-glyph boxes in the PDF,
    replace a few common icons used in this project with plain text labels.
    """
    if not s:
        return s

    # Remove variation selectors / joiners that commonly appear in emoji sequences.
    s = s.replace("\ufe0f", "").replace("\u200d", "")

    # If we can render emoji via an embedded emoji-capable TTF, keep symbols.
    if emoji_font_available:
        return s

    # Replace commonly used icons with text labels (keeps meaning on all fonts).
    replacements = {
        "🎤": "Singer: ",
        "✍": "Lyrics: ",
        "🎼": "Music: ",
    }
    for k, v in replacements.items():
        s = s.replace(k, v)
    return s


def _register_emoji_font(base_dir: Path, explicit_path: str = "") -> Optional[str]:
    """
    Register a font that can render a few emoji/symbols (🎤 ✍ 🎼) as monochrome glyphs.
    Note: macOS Apple Color Emoji (.ttc) generally won't work with ReportLab TTFont.
    Best approach: provide a .ttf in ./fonts (e.g., Symbola.ttf, NotoSansSymbols2-Regular.ttf).
    """
    candidates: List[Path] = []
    if explicit_path:
        p = safe_path(base_dir, explicit_path)
        if p:
            candidates.append(p)

    fonts_dir = (base_dir / "fonts")
    if fonts_dir.exists():
        for pat in ("*Symbola*.ttf", "*NotoSansSymbols2*.ttf", "*Segoe*Symbol*.ttf", "*Emoji*.ttf"):
            candidates.extend(sorted(fonts_dir.glob(pat)))

    # Common system locations (best-effort)
    candidates.extend([
        Path("/usr/share/fonts/truetype/noto/NotoSansSymbols2-Regular.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansSymbols2-Regular.ttf"),
        Path("/Library/Fonts/Symbola.ttf"),
        Path("/System/Library/Fonts/Symbola.ttf"),
        Path("C:/Windows/Fonts/seguisym.ttf"),
    ])

    for fp in candidates:
        try:
            if not fp.exists() or not fp.is_file():
                continue
            font_name = f"SongbookEmoji-{fp.stem}"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(fp)))
            return font_name
        except Exception:
            continue
    return None


def load_songbook(json_path: Path) -> Tuple[str, Dict[str, Any], List[Dict[str, Any]]]:
    data = json.loads(json_path.read_text(encoding="utf-8"))

    book_title = data.get("book_title", "My Songbook")
    book_meta = data.get("book_meta", {}) or {}

    songs = data.get("songs", [])
    if not isinstance(songs, list):
        raise ValueError("songs.json: 'songs' must be a list")

    seen = set()
    for s in songs:
        if "id" not in s or "title" not in s:
            raise ValueError("Each song must have 'id' and 'title'")
        if s["id"] in seen:
            raise ValueError(f"Duplicate song id: {s['id']}")
        seen.add(s["id"])

        # Default: export songs unless explicitly disabled.
        s.setdefault("export", True)
        s["export"] = _boolish(s.get("export", True), default=True)

        s.setdefault("info", [])
        s.setdefault("sections", [])
        s.setdefault("thumbnail", "")
        s.setdefault("background", "")
        # Optional: per-song override for how background is rendered in PDF.
        # Values: "cover" (fill/crop), "tile" (repeat), "contain" (fit inside)
        s.setdefault("background_mode", "")

    songs_export = [s for s in songs if _boolish(s.get("export", True), default=True)]
    return book_title, book_meta, songs_export


# -----------------------------
# Tokenization fallback (if tokens missing)
# -----------------------------
def tokenize_fallback(indian_line: str) -> List[str]:
    """
    Minimal fallback tokenization if tokens[] not provided.
    Accepts basic:
      S R G m M P D N, r g d n, with optional:
      - meend: G~R
      - kan: (R)G
      - hold markers: ... or … or -- or :
      - octaves: S' or S’ (high), S. (low). Also accepts legacy low: ,S
      - komal inline: D(k), N(k)
    """
    s = indian_line.strip()
    if not s:
        return []

    # Normalize separators
    s = s.replace("|", " ").replace("…", " ... ")
    s = re.sub(r"\s+", " ", s)

    parts = s.split(" ")
    out: List[str] = []

    # patterns
    komal_inline = re.compile(r"^([RGDN])\((k|K)\)$")  # D(k) -> d
    note_simple = re.compile(r"^,?[SRGmMPDNrgdn](?:[’']|\.|)?(?::)?$")  # + low + octave + hold(:)
    meend = re.compile(r"^([SRGmMPDNrgdn])~([SRGmMPDNrgdn])$")
    kan = re.compile(r"^\(([SRGmMPDNrgdn])\)([SRGmMPDNrgdn])([’']|\.|)?$")

    def norm_note(tok: str) -> str:
        tok = tok.strip()
        low = tok.startswith(",")
        if low:
            tok = tok[1:]

        octv = ""
        if tok.endswith(("’", "'")):
            octv = "'"
            tok = tok[:-1]
        elif tok.endswith("."):
            octv = "."
            tok = tok[:-1]

        m = komal_inline.match(tok)
        if m:
            tok = m.group(1).lower()

        tok = tok.replace("(k)", "").replace("(K)", "")

        if low:
            return "," + tok + octv
        return tok + octv

    for p in parts:
        p = p.strip()
        if not p:
            continue

        # detect hold patterns like G..., G---, etc -> canonical ":"
        hold = ""
        if "..." in p or re.search(r"--+$", p) or re.search(r"\.{3,}$", p) or p.endswith(":"):
            hold = ":"
            p = re.sub(r":$", "", p)
            p = re.sub(r"--+$", "", p)
            p = re.sub(r"\.{3,}$", "", p)
            p = p.replace("...", "")
            p = p.strip()

        km = kan.match(p)
        if km:
            out.append(f"({norm_note(km.group(1))}){norm_note(km.group(2) + (km.group(3) or ''))}")
            continue

        mm = meend.match(p)
        if mm:
            out.append(f"{norm_note(mm.group(1))}~{norm_note(mm.group(2))}")
            continue

        if komal_inline.match(p):
            out.append(norm_note(p) + hold)
            continue

        if note_simple.match(p):
            out.append(norm_note(p) + hold)
            continue

    return out


# -----------------------------
# Derivation helpers (word-style Indian -> tokens / western)
# -----------------------------
_RE_KOMAL_INLINE_ANYWHERE = re.compile(r"([RGDN])\((k|K)\)")  # R(k) -> r (not word-boundary)


def _indian_words_to_letter_notation(indian_line: str) -> str:
    """
    Convert display-style Indian notation (Sa/Re/Ga...) into letter-token notation (S/R/G...).

    This lets us keep songs.json clean (no stored tokens/western) while still generating:
    - Western display lines
    - MusicXML

    Examples:
      "Ni Sa DhaDhaGa Dha Pa" -> "N S D D G D P"
      "Pa Pa ,PaNiSaGa Re ,Pa Ni SaGaRe~Sa" -> "P P ,P N S G R ,P N S G R~S"
      "Pa Ni Re Sa: Sa: Sa: Sa:" -> "P N R S: S: S: S:"
    """
    s = (indian_line or "").strip()
    if not s:
        return ""

    # Normalize curly apostrophe
    s = s.replace("’", "'")

    # Accept user-friendly display ordering where octave marker appears before (k)/(T),
    # e.g. Re.(k), Re'(k), Ma'(T). Internally we normalize to Re(k)., Re(k)', Ma(T)'.
    # Do this for both word-style (Re...) and letter-style (R...) inputs.
    s = re.sub(r"(?i)\b(Sa|Re|Ga|Ma|Pa|Dha|Ni)([.'])(\((?:k|K|t|T)\))", r"\1\3\2", s)
    s = re.sub(r"(?i)([SRGmMPDNrgdn])([.'])(\((?:k|K|t|T)\))", r"\1\3\2", s)

    # Tivra Ma forms: Ma(T) / M(T) / M#
    s = re.sub(r"Ma\((?:T|t)\)", "M", s, flags=re.IGNORECASE)
    s = re.sub(r"\bMa#", "M", s, flags=re.IGNORECASE)
    s = re.sub(r"\bM\((?:T|t)\)", "M", s)
    s = re.sub(r"\bM#", "M", s)

    # Komal full-word forms
    s = re.sub(r"Re\((?:k|K)\)", "r", s, flags=re.IGNORECASE)
    s = re.sub(r"Ga\((?:k|K)\)", "g", s, flags=re.IGNORECASE)
    s = re.sub(r"Dha\((?:k|K)\)", "d", s, flags=re.IGNORECASE)
    s = re.sub(r"Ni\((?:k|K)\)", "n", s, flags=re.IGNORECASE)

    # Komal inline letter forms (R(k), etc.) anywhere in the string
    s = _RE_KOMAL_INLINE_ANYWHERE.sub(lambda m: m.group(1).lower(), s)

    # Basic swara words → letters (match even when concatenated: "SaGaRe")
    # Order matters: Dha before others, and Ma last (after Ma(T) handled above).
    replacements = [
        ("Dha", "D"),
        ("Ni", "N"),
        ("Sa", "S"),
        ("Re", "R"),
        ("Ga", "G"),
        ("Pa", "P"),
        ("Ma", "m"),
    ]
    for w, rep in replacements:
        s = re.sub(w, rep, s, flags=re.IGNORECASE)

    # Insert spaces between adjacent note symbols so tokenization can split reliably.
    # Examples: "DDG" -> "D D G", "S'G" -> "S' G", "SGR~S" -> "S G R~S"
    s = re.sub(r"([SRGmMPDNrgdn](?:[']|\.)?)(?=[SRGmMPDNrgdn])", r"\1 ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def tokenize_from_indian(indian_line: str) -> List[str]:
    """
    Primary tokenizer for this project.

    Supports:
    - letter style: "S R G m P D N"
    - word style: "Sa Re Ga Ma Pa Dha Ni" (including concatenations like "SaGaRe")
    """
    toks = tokenize_fallback(indian_line)
    if toks:
        return toks
    converted = _indian_words_to_letter_notation(indian_line)
    return tokenize_fallback(converted)


def western_from_indian(indian_line: str, *, include_octave: bool = False, default_octave: int = 4) -> str:
    """
    Generate a compact Western display string from an Indian display line.

    - SA=C mapping (same as MusicXML generator).
    - Preserves: meend "~", kan "(x)y", hold ":" and comma-prefix markers.
    - By default omits octave numbers (to match the existing songs.json style).
    """
    tokens = tokenize_from_indian(indian_line)
    if not tokens:
        return ""

    out: List[str] = []
    kan_re = re.compile(r"^\(([^)]+)\)(.+)$")
    meend_re = re.compile(r"^(.+)~(.+)$")

    def western_note_label(tok: str) -> str:
        pn = _parse_note_token(tok, default_octave=default_octave)
        if not pn:
            return ""
        if include_octave:
            return pn.western_label
        acc = "♭" if pn.alter == -1 else ("#" if pn.alter == 1 else "")
        return f"{pn.step}{acc}"

    for tok0 in tokens:
        t = tok0.strip()
        if not t:
            continue

        comma_prefix = t.startswith(",")
        if comma_prefix:
            t = t[1:]

        hold = t.endswith(":")
        if hold:
            t = t[:-1]

        seg = ""
        km = kan_re.match(t)
        if km:
            seg = f"({western_note_label(km.group(1))}){western_note_label(km.group(2))}"
        else:
            mm = meend_re.match(t)
            if mm:
                seg = f"{western_note_label(mm.group(1))}~{western_note_label(mm.group(2))}"
            else:
                seg = western_note_label(t)

        if not seg:
            continue
        if hold:
            seg += ":"

        if comma_prefix:
            if out:
                out[-1] = out[-1].rstrip() + ","
            else:
                seg = "," + seg

        out.append(seg)

    return " ".join(out).strip()


# -----------------------------
# PDF generator
# -----------------------------
def make_pdf(
    book_title: str,
    book_meta: Dict[str, Any],
    songs: List[Dict[str, Any]],
    base_dir: Path,
    out_path: Path,
    bg_opacity: float = 0.08,
    page: str = "A5",
    pdf_bg_mode: str = "cover",
    pdf_emoji_font: str = "",
    *,
    show_indian: bool = True,
    show_western: bool = True,
) -> None:
    _require_reportlab()
    pagesize = A5 if page.upper() == "A5" else letter
    W, H = pagesize
    c = canvas.Canvas(str(out_path), pagesize=pagesize)

    margin = 0.50 * inch
    # Use a smaller bottom margin so the body can fill more of the page.
    bottom_margin = 0.18 * inch
    back_link_y = 0.16 * inch  # y-position for "Back to Index" (closer to bottom)
    text_fonts = _register_unicode_font_pack(base_dir)
    unicode_font_available = text_fonts.unicode_ok
    emoji_font = _register_emoji_font(base_dir, explicit_path=pdf_emoji_font)
    emoji_font_available = emoji_font is not None

    def draw_text_with_emoji_fallback(x: float, y: float, text: str, *, base_font: str, base_size: float) -> None:
        """
        Draws text but renders specific emoji glyphs using emoji_font (if available).
        This is a very small, targeted fallback (we only care about a few icons).
        """
        t = _pdf_safe_text(text, emoji_font_available=emoji_font_available)

        # If no emoji font, draw normally.
        if not emoji_font_available:
            c.setFont(base_font, base_size)
            c.drawString(x, y, t)
            return

        # With emoji font: render emoji chars with emoji font, everything else with base font.
        emoji_chars = {"🎤", "✍", "🎼"}
        c.setFont(base_font, base_size)
        cur_x = x
        buf: List[str] = []

        def flush_buf() -> None:
            nonlocal cur_x, buf
            if not buf:
                return
            seg = "".join(buf)
            c.setFont(base_font, base_size)
            c.drawString(cur_x, y, seg)
            cur_x += pdfmetrics.stringWidth(seg, base_font, base_size)
            buf = []

        for ch in t:
            if ch in emoji_chars:
                flush_buf()
                c.setFont(emoji_font, base_size)
                c.drawString(cur_x, y, ch)
                cur_x += pdfmetrics.stringWidth(ch, emoji_font, base_size)
                c.setFont(base_font, base_size)
            else:
                buf.append(ch)
        flush_buf()

    def draw_bg(img_path: Optional[Path], *, mode: str) -> None:
        if not img_path:
            return
        try:
            c.saveState()
            if hasattr(c, "setFillAlpha"):
                c.setFillAlpha(bg_opacity)

            img = ImageReader(str(img_path))
            iw, ih = img.getSize()
            if not iw or not ih:
                c.restoreState()
                return

            mode = (mode or "cover").strip().lower()

            if mode == "tile":
                # Scale each tile to "contain" once, then repeat to cover the page.
                scale = min(W / iw, H / ih)
                tw = max(1.0, iw * scale)
                th = max(1.0, ih * scale)
                nx = int(math.ceil(W / tw))
                ny = int(math.ceil(H / th))
                for ix in range(nx):
                    for iy in range(ny):
                        c.drawImage(
                            img,
                            ix * tw,
                            iy * th,
                            width=tw,
                            height=th,
                            preserveAspectRatio=False,
                            mask="auto",
                        )
            elif mode == "contain":
                # Fit entire image within page; center it.
                scale = min(W / iw, H / ih)
                dw = iw * scale
                dh = ih * scale
                x = (W - dw) / 2.0
                y = (H - dh) / 2.0
                c.drawImage(
                    img,
                    x,
                    y,
                    width=dw,
                    height=dh,
                    preserveAspectRatio=False,
                    mask="auto",
                )
            else:
                # Default "cover": fill the page, crop overflow; center.
                scale = max(W / iw, H / ih)
                dw = iw * scale
                dh = ih * scale
                x = (W - dw) / 2.0
                y = (H - dh) / 2.0
                c.drawImage(
                    img,
                    x,
                    y,
                    width=dw,
                    height=dh,
                    preserveAspectRatio=False,
                    mask="auto",
                )
            c.restoreState()
        except Exception:
            pass

    def draw_header(song: Dict[str, Any]) -> float:
        thumb = safe_path(base_dir, song.get("thumbnail", ""))
        y = H - margin
        x = margin
        thumb_w = 1.0 * inch
        thumb_h = 1.0 * inch

        has_thumb = False
        if thumb:
            try:
                c.drawImage(str(thumb), x, y - thumb_h, width=thumb_w, height=thumb_h, mask="auto")
                has_thumb = True
            except Exception:
                has_thumb = False

        if not has_thumb:
            c.rect(x, y - thumb_h, thumb_w, thumb_h, stroke=1, fill=0)
            c.setFont(text_fonts.italic, 8.5)
            c.drawCentredString(x + thumb_w / 2, y - thumb_h / 2, "Thumbnail")

        tx = x + thumb_w + 0.25 * inch
        c.setFont(text_fonts.bold, 14)
        c.drawString(tx, y - 0.20 * inch, song["title"])

        c.setFont(text_fonts.regular, 9.5)
        iy = y - 0.42 * inch
        for line in (song.get("info", []) or [])[:10]:
            draw_text_with_emoji_fallback(tx, iy, str(line), base_font=text_fonts.regular, base_size=9.5)
            iy -= 0.16 * inch

        return y - 1.25 * inch

    def wrap_text(text: str, font_name: str, font_size: float, max_width: float) -> List[str]:
        """
        Very small word-wrap helper for ReportLab drawString.
        """
        t = (text or "").strip()
        if not t:
            return [""]
        words = t.split(" ")
        lines: List[str] = []
        cur: List[str] = []

        for w0 in words:
            w = w0
            tentative = (" ".join(cur + [w])).strip()
            if pdfmetrics.stringWidth(tentative, font_name, font_size) <= max_width or not cur:
                cur.append(w)
            else:
                lines.append(" ".join(cur).rstrip())
                cur = [w]
        if cur:
            lines.append(" ".join(cur).rstrip())
        return lines or [t]

    def draw_wrapped_lines(
        x: float,
        y: float,
        text: str,
        *,
        font_name: str,
        font_size: float,
        leading: float,
        max_width: float,
        italic_fallback: bool = False,
    ) -> float:
        """
        Draw wrapped text at (x,y) top-down, returns new y.
        """
        t = _western_pdf_safe(text, unicode_font_available=unicode_font_available) if italic_fallback else text
        lines = wrap_text(t, font_name, font_size, max_width)
        c.setFont(font_name, font_size)
        for ln in lines:
            c.drawString(x, y, ln)
            y -= leading
        return y

    def measure_block_height(lyric: str, indian: str, western: str, *, col_width: float) -> float:
        """
        Estimate height (points) required for this block in a column.
        """
        # Font metrics (points)
        # Styling rules:
        # - Section headers (INTRO/MUKHDA/ANTARA/INTERLUDE) are handled elsewhere (bigger/bold).
        # - Lyrics should be bold but similar size to notations.
        ly_font, ly_size, ly_lead = text_fonts.bold, 9.0, 11.0
        in_font, in_size, in_lead = text_fonts.regular, 9.0, 11.0
        we_font, we_size, we_lead = text_fonts.regular, 9.0, 11.0

        ly_lines = wrap_text(lyric, ly_font, ly_size, col_width)
        in_lines = wrap_text(indian, in_font, in_size, col_width) if show_indian else []
        we_lines = (
            wrap_text(_western_pdf_safe(western, unicode_font_available=unicode_font_available), we_font, we_size, col_width)
            if show_western
            else []
        )

        # Spacing between rows + block gap (tighter if a row is omitted)
        h = 0.0
        h += len(ly_lines) * ly_lead
        if show_indian and in_lines:
            h += 3.0 + len(in_lines) * in_lead
        if show_western and we_lines:
            h += (2.0 if (show_indian and in_lines) else 3.0) + len(we_lines) * we_lead
        # Block gap
        h += 7.0
        return h

    def draw_line_block_in_column(
        lyric: str,
        indian: str,
        western: str,
        *,
        x: float,
        y: float,
        col_width: float,
    ) -> float:
        """
        Draw a 3-line (lyrics/indian/western) block in a column with wrapping.
        Returns the new y.
        """
        # Lyrics
        y = draw_wrapped_lines(
            x,
            y,
            lyric,
            font_name=text_fonts.bold,
            font_size=9.0,
            leading=11.0,
            max_width=col_width,
        )
        # Optional notations
        if show_indian:
            y -= 3.0
            y = draw_wrapped_lines(
                x,
                y,
                indian,
                font_name=text_fonts.regular,
                font_size=9.0,
                leading=11.0,
                max_width=col_width,
            )

        if show_western:
            y -= 2.0 if show_indian else 3.0
            y = draw_wrapped_lines(
                x,
                y,
                western,
                font_name=text_fonts.regular,
                font_size=9.0,
                leading=11.0,
                max_width=col_width,
                italic_fallback=True,  # if font isn't Unicode-capable, maps ♭/♯ to b/# for readability
            )

        y -= 7.0
        return y

    # Index background:
    # Prefer book_meta.cover_image if provided; fall back to first song background.
    first_bg = None
    cover_rel = str((book_meta.get("cover_image") or "")).strip()
    if cover_rel:
        first_bg = safe_path(base_dir, cover_rel)
    if not first_bg:
        for s in songs:
            bg = safe_path(base_dir, s.get("background", ""))
            if bg:
                first_bg = bg
                break
    draw_bg(first_bg, mode=pdf_bg_mode)

    # Index page
    c.bookmarkPage("dest_index")
    c.setFont(text_fonts.bold, 18)
    c.drawCentredString(W / 2, H - 1.0 * inch, book_title)
    # c.setFont("Times-Roman", 10)
    # c.drawCentredString(W / 2, H - 1.2 * inch, "Mobile-friendly single-column layout")
    c.setFont(text_fonts.bold, 12)
    c.drawString(margin, H - 1.65 * inch, "Index")

    y = H - 1.95 * inch
    for s in songs:
        dest = f"dest_{s['id']}"
        c.setFont(text_fonts.regular, 12)
        c.drawString(margin + 0.10 * inch, y, f"• {s['title']}")
        c.linkRect("", dest, (margin, y - 0.05 * inch, W - margin, y + 0.18 * inch), relative=0, thickness=0)
        y -= 0.32 * inch
        if y < margin + 0.8 * inch:
            c.showPage()
            draw_bg(first_bg, mode=pdf_bg_mode)
            c.bookmarkPage("dest_index_cont")
            y = H - margin

    c.showPage()

    # Song pages
    for s in songs:
        bg = safe_path(base_dir, s.get("background", ""))
        per_song_mode = (s.get("background_mode") or "").strip()
        dest = f"dest_{s['id']}"

        def start_song_page(*, cont: bool) -> Tuple[float, float, float, float]:
            """
            Returns (body_top_y, col1_x, col2_x, col_width).
            """
            draw_bg(bg, mode=per_song_mode or pdf_bg_mode)
            if not cont:
                c.bookmarkPage(dest)
            title = s["title"] + (" (cont.)" if cont else "")
            y0 = draw_header({"title": title, "info": s.get("info", []), "thumbnail": s.get("thumbnail", "")})

            # Header/body separator label
            c.setFont(text_fonts.bold, 9.5)
            y0 -= 0.26 * inch

            gutter = 0.30 * inch
            col_width = (W - 2 * margin - gutter)
            col_width = col_width / 2.0
            col1_x = margin
            col2_x = margin + col_width + gutter
            return (y0, col1_x, col2_x, col_width)

        body_top, col1_x, col2_x, col_w = start_song_page(cont=False)
        col_y = body_top
        col_idx = 0  # 0=left, 1=right

        def advance_column_or_page(required_height: float) -> None:
            """
            Ensure there is vertical space in current column; otherwise switch column or page.
            """
            nonlocal body_top, col1_x, col2_x, col_w, col_y, col_idx
            # Reserve only a tight strip for the "Back to Index" footer.
            min_y = back_link_y + 0.18 * inch
            if col_y - required_height >= min_y:
                return
            if col_idx == 0:
                col_idx = 1
                col_y = body_top
                if col_y - required_height >= min_y:
                    return
            # New page (keep header, background)
            c.showPage()
            body_top, col1_x, col2_x, col_w = start_song_page(cont=True)
            col_idx = 0
            col_y = body_top

        sections = s.get("sections", []) or []
        for section in sections:
            sec_name = (section.get("name") or "").strip()
            if sec_name:
                # Section heading in-column (wrap-aware)
                heading = sec_name.upper()
                heading_font_size = 11.5
                heading_leading = 13.5
                h = len(wrap_text(heading, text_fonts.bold, heading_font_size, col_w)) * heading_leading + 7.0
                advance_column_or_page(h)
                x = col1_x if col_idx == 0 else col2_x
                c.setFont(text_fonts.bold, heading_font_size)
                for ln in wrap_text(heading, text_fonts.bold, heading_font_size, col_w):
                    c.drawString(x, col_y, ln)
                    col_y -= heading_leading
                col_y -= 7.0

            for line in (section.get("lines") or []):
                lyric = str(line.get("lyrics", "") or "")
                indian = str(line.get("indian", "") or "")
                western = str(line.get("western", "") or "")
                if show_western and not western.strip():
                    western = western_from_indian(indian)

                if not lyric and not indian and not western:
                    continue

                needed = measure_block_height(lyric, indian, western, col_width=col_w)
                advance_column_or_page(needed)

                x = col1_x if col_idx == 0 else col2_x
                col_y = draw_line_block_in_column(lyric, indian, western, x=x, y=col_y, col_width=col_w)

        c.setFont(text_fonts.italic, 9)
        c.drawString(margin, back_link_y, "Back to Index")
        c.linkRect(
            "",
            "dest_index",
            (margin, back_link_y - 0.10 * inch, margin + 1.45 * inch, back_link_y + 0.11 * inch),
            relative=0,
            thickness=0,
        )
        c.showPage()

    c.save()


# -----------------------------
# EPUB generator (EPUB2)
# -----------------------------
def make_epub(
    book_title: str,
    book_meta: Dict[str, Any],
    songs: List[Dict[str, Any]],
    base_dir: Path,
    out_path: Path,
    bg_opacity: float = 0.10,
) -> None:
    tmp = out_path.parent / "_epub_tmp"
    if tmp.exists():
        shutil.rmtree(tmp)

    (tmp / "META-INF").mkdir(parents=True)
    (tmp / "OEBPS" / "images").mkdir(parents=True)

    # Required files
    (tmp / "mimetype").write_text("application/epub+zip", encoding="utf-8")
    (tmp / "META-INF" / "container.xml").write_text(
        """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""",
        encoding="utf-8",
    )

    css = f"""
body {{ font-family: serif; line-height: 1.45; margin: 0; padding: 0; }}
.page {{ padding: 1.0rem 1.0rem 1.5rem 1.0rem; position: relative; }}
.bg {{ position: fixed; inset: 0; background-size: cover; background-position: center;
      opacity: {bg_opacity}; z-index: 0; }}
.content {{ position: relative; z-index: 1; }}
.header {{ display: flex; gap: 0.9rem; align-items: flex-start; margin-bottom: 0.9rem; }}
.thumb {{ width: 96px; height: 96px; object-fit: cover; border-radius: 10px; }}
.title {{ font-size: 1.4rem; font-weight: 700; margin: 0; }}
.meta {{ font-size: 0.95rem; margin-top: 0.25rem; }}
.block {{ margin: 0.75rem 0 0.9rem 0; }}
.ly {{ font-weight: 700; margin: 0 0 0.25rem 0; }}
.sa {{ font-size: 1.05rem; margin: 0 0 0.15rem 0; }}
.we {{ font-size: 0.95rem; font-style: italic; margin: 0; opacity: 0.9; }}
a {{ text-decoration: none; }}
.index a {{ display: block; padding: 0.45rem 0; }}
.small {{ font-size: 0.9rem; opacity: 0.85; }}
"""
    (tmp / "OEBPS" / "styles.css").write_text(css, encoding="utf-8")

    def xhtml_doc(title: str, body: str) -> str:
        return f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
  <title>{xesc(title)}</title>
  <meta charset="utf-8"/>
  <link rel="stylesheet" type="text/css" href="styles.css"/>
</head>
<body>
{body}
</body>
</html>
"""

    def copy_image(rel_path: str) -> Optional[str]:
        p = safe_path(base_dir, rel_path)
        if not p:
            return None
        dest = tmp / "OEBPS" / "images" / p.name
        shutil.copy(p, dest)
        return f"images/{p.name}"

    # Optional cover
    cover_rel = (book_meta.get("cover_image") or "").strip()
    cover_epub = copy_image(cover_rel) if cover_rel else None

    # Index page
    index_items = "\n".join([f'<a href="{s["id"]}.xhtml">{xesc(s["title"])}</a>' for s in songs])
    index_body = f"""
<div class="page">
  <div class="content">
    <h1 class="title">{xesc(book_title)}</h1>
    <div class="meta small">Tap a song. EPUB is reflowable (fonts scale).</div>
    <div class="index" style="margin-top:1rem;">
      {index_items}
    </div>
  </div>
</div>
"""
    (tmp / "OEBPS" / "index.xhtml").write_text(xhtml_doc("Index", index_body), encoding="utf-8")

    # Cover page (if cover image exists)
    if cover_epub:
        cover_body = f"""
<div class="page">
  <div class="content" style="text-align:center;">
    <h1 class="title">{xesc(book_title)}</h1>
    <div style="margin-top:1rem;">
      <img src="{cover_epub}" alt="cover" style="max-width:90%; border-radius:14px;"/>
    </div>
  </div>
</div>
"""
        (tmp / "OEBPS" / "cover.xhtml").write_text(xhtml_doc("Cover", cover_body), encoding="utf-8")

    # Song pages
    for s in songs:
        thumb = copy_image(s.get("thumbnail", ""))
        bg = copy_image(s.get("background", ""))

        bg_div = f'<div class="bg" style="background-image:url({bg});"></div>' if bg else ""
        thumb_img = f'<img class="thumb" src="{thumb}" alt="thumbnail"/>' if thumb else ""

        meta_html = "<br/>".join([xesc(str(line)) for line in (s.get("info", []) or [])[:12]])

        blocks: List[str] = []
        for section in (s.get("sections") or []):
            sec_name = (section.get("name") or "").strip()
            if sec_name:
                blocks.append(f'<div class="block"><p class="ly">{xesc(sec_name.upper())}</p></div>')
            for line in (section.get("lines") or []):
                lyric = xesc(str(line.get("lyrics", "") or ""))
                indian_raw = str(line.get("indian", "") or "")
                indian = xesc(indian_raw)
                western_raw = str(line.get("western", "") or "")
                if not western_raw.strip():
                    western_raw = western_from_indian(indian_raw)
                western = xesc(western_raw)
                if not (lyric or indian or western):
                    continue
                blocks.append(f"""
<div class="block">
  <p class="ly">{lyric}</p>
  <p class="sa">{indian}</p>
  <p class="we">{western}</p>
</div>""")

        song_body = f"""
<div class="page">
  {bg_div}
  <div class="content">
    <div class="header">
      {thumb_img}
      <div>
        <h1 class="title">{xesc(s["title"])}</h1>
        <div class="meta">{meta_html}</div>
      </div>
    </div>
    {''.join(blocks)}
    <div class="small"><a href="index.xhtml">Back to Index</a></div>
  </div>
</div>
"""
        (tmp / "OEBPS" / f'{s["id"]}.xhtml').write_text(xhtml_doc(s["title"], song_body), encoding="utf-8")

    # OPF + NCX
    creator = str(book_meta.get("creator", "") or "").strip()
    publisher = str(book_meta.get("publisher", "") or "").strip()
    language = str(book_meta.get("language", "en") or "en").strip()
    isbn = str(book_meta.get("isbn", "") or "").strip()
    identifier = isbn if isbn else f"urn:uuid:songbook-{date.today().isoformat()}"

    manifest: List[str] = [
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
        '<item id="css" href="styles.css" media-type="text/css"/>',
        '<item id="index" href="index.xhtml" media-type="application/xhtml+xml"/>',
    ]
    spine: List[str] = []
    navpoints: List[str] = []

    play = 1

    # Cover first in spine/TOC if present
    if cover_epub:
        manifest.append('<item id="coverpage" href="cover.xhtml" media-type="application/xhtml+xml"/>')
        manifest.append(f'<item id="coverimage" href="{cover_epub}" media-type="image/png"/>')
        spine.append('<itemref idref="coverpage"/>')
        navpoints.append(f'<navPoint id="nav{play}" playOrder="{play}"><navLabel><text>Cover</text></navLabel><content src="cover.xhtml"/></navPoint>')
        play += 1

    spine.append('<itemref idref="index"/>')
    navpoints.append(f'<navPoint id="nav{play}" playOrder="{play}"><navLabel><text>Index</text></navLabel><content src="index.xhtml"/></navPoint>')
    play += 1

    for s in songs:
        manifest.append(f'<item id="{s["id"]}" href="{s["id"]}.xhtml" media-type="application/xhtml+xml"/>')
        spine.append(f'<itemref idref="{s["id"]}"/>')
        navpoints.append(f'<navPoint id="nav{play}" playOrder="{play}"><navLabel><text>{xesc(s["title"])}</text></navLabel><content src="{s["id"]}.xhtml"/></navPoint>')
        play += 1

    # Include images in manifest
    img_files = list((tmp / "OEBPS" / "images").glob("*"))
    for img in img_files:
        ext = img.suffix.lower()
        mt = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        # avoid duplicate id collisions
        item_id = slugify(img.stem)
        manifest.append(f'<item id="img-{item_id}" href="images/{img.name}" media-type="{mt}"/>')

    # Optional EPUB2 cover metadata
    cover_meta = ""
    if cover_epub:
        cover_meta = '<meta name="cover" content="coverimage"/>'

    content_opf = f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>{xesc(book_title)}</dc:title>
    <dc:language>{xesc(language)}</dc:language>
    <dc:identifier id="BookId">{xesc(identifier)}</dc:identifier>
    {f"<dc:creator>{xesc(creator)}</dc:creator>" if creator else ""}
    {f"<dc:publisher>{xesc(publisher)}</dc:publisher>" if publisher else ""}
    <dc:date>{date.today().isoformat()}</dc:date>
    {cover_meta}
  </metadata>
  <manifest>
    {''.join(manifest)}
  </manifest>
  <spine toc="ncx">
    {''.join(spine)}
  </spine>
</package>
"""
    (tmp / "OEBPS" / "content.opf").write_text(content_opf, encoding="utf-8")

    toc_ncx = f"""<?xml version="1.0" encoding="UTF-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="{xesc(identifier)}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{xesc(book_title)}</text></docTitle>
  <navMap>
    {''.join(navpoints)}
  </navMap>
</ncx>
"""
    (tmp / "OEBPS" / "toc.ncx").write_text(toc_ncx, encoding="utf-8")

    # Zip to EPUB: mimetype must be first and stored
    if out_path.exists():
        out_path.unlink()

    with zipfile.ZipFile(out_path, "w") as z:
        z.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        for root, _, files in os.walk(tmp):
            for fn in files:
                p = Path(root) / fn
                rel = p.relative_to(tmp).as_posix()
                if rel == "mimetype":
                    continue
                z.write(p, rel, compress_type=zipfile.ZIP_DEFLATED)

    shutil.rmtree(tmp)


# -----------------------------
# DOCX generator
# -----------------------------
def make_docx(
    book_title: str,
    songs: List[Dict[str, Any]],
    base_dir: Path,
    out_path: Path,
) -> None:
    """
    Create a Word document from songs.json:
    - Header: thumbnail + title + info
    - Body: 2-column table (lyrics + indian + western)
    """
    _require_docx()

    doc = Document()  # type: ignore[misc]

    # Title page
    doc.add_heading(book_title, level=0)
    doc.add_paragraph("Generated from songs.json")
    doc.add_page_break()

    def set_run_font(run, *, bold: bool, size_pt: float) -> None:
        run.bold = bold
        run.font.size = Pt(size_pt)  # type: ignore[misc]

    def add_block(cell, lyric: str, indian: str, western: str) -> None:
        # Lyrics (slightly larger)
        p1 = cell.add_paragraph()
        r1 = p1.add_run(lyric)
        set_run_font(r1, bold=True, size_pt=10)

        # Indian (one size smaller)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(indian)
        set_run_font(r2, bold=False, size_pt=9)

        # Western (one size smaller, normal)
        p3 = cell.add_paragraph()
        r3 = p3.add_run(western)
        set_run_font(r3, bold=False, size_pt=9)

    for idx, s in enumerate(songs):
        if idx != 0:
            doc.add_page_break()

        # Header table: thumbnail | info
        header = doc.add_table(rows=1, cols=2)
        header.autofit = True
        left = header.cell(0, 0)
        right = header.cell(0, 1)

        # Thumbnail
        thumb = safe_path(base_dir, s.get("thumbnail", ""))
        if thumb:
            try:
                left.paragraphs[0].add_run().add_picture(str(thumb), width=Inches(1.2))  # type: ignore[misc]
            except Exception:
                left.text = "Thumbnail"
        else:
            left.text = "Thumbnail"

        # Title + info
        ptitle = right.paragraphs[0]
        rtitle = ptitle.add_run(str(s.get("title", "")))
        set_run_font(rtitle, bold=True, size_pt=14)

        for info_line in (s.get("info", []) or [])[:12]:
            p = right.add_paragraph(str(info_line))
            if p.runs:
                set_run_font(p.runs[0], bold=False, size_pt=9)

        doc.add_paragraph("")  # spacer

        # Body: 2-column table for lines
        body = doc.add_table(rows=0, cols=2)
        body.autofit = True

        # Build a flat list of display blocks including section headings.
        blocks: List[Tuple[str, str, str, bool]] = []
        for section in (s.get("sections") or []):
            sec_name = (section.get("name") or "").strip()
            if sec_name:
                blocks.append((sec_name.upper(), "", "", True))
            for line in (section.get("lines") or []):
                lyric = str(line.get("lyrics", "") or "").strip()
                indian = str(line.get("indian", "") or "").strip()
                western = str(line.get("western", "") or "").strip()
                if not western:
                    western = western_from_indian(indian)
                if not (lyric or indian or western):
                    continue
                blocks.append((lyric, indian, western, False))

        i = 0
        while i < len(blocks):
            ly, ind, we, is_heading = blocks[i]
            if is_heading:
                row = body.add_row()
                merged = row.cells[0].merge(row.cells[1])
                p = merged.paragraphs[0]
                r = p.add_run(ly)
                set_run_font(r, bold=True, size_pt=10)
                i += 1
                continue

            row = body.add_row()
            # left cell
            add_block(row.cells[0], ly, ind, we)
            # right cell (next block if available and not a heading)
            if i + 1 < len(blocks) and not blocks[i + 1][3]:
                ly2, ind2, we2, _ = blocks[i + 1]
                add_block(row.cells[1], ly2, ind2, we2)
                i += 2
            else:
                row.cells[1].text = ""
                i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# -----------------------------
# MusicXML generator (per song)
# -----------------------------
@dataclass
class ParsedNote:
    step: str
    alter: int
    octave: int
    indian_label: str
    western_label: str


# Pitch mapping is loaded from notation_mapping.json into _TOKEN_TO_WESTERN (module-level).

# Ornament patterns
_RE_MEEND = re.compile(r"^(.+)~(.+)$")
_RE_KAN = re.compile(r"^\(([^)]+)\)(.+)$")  # (R)G -> grace R + note G


def _parse_note_token(tok: str, default_octave: int = 4) -> Optional[ParsedNote]:
    """
    Token examples supported:
      S R G m M P D N
      r g d n  (komal)
      D(k) N(k) etc are expected to be normalized earlier; fallback handled below
      high: S' or S’   (apostrophe)
      low: ,S   or S.  (dot)
    """
    raw = tok.strip()
    if not raw:
        return None

    # normalize curly apostrophe
    raw = raw.replace("’", "'")

    low = False
    if raw.startswith(","):
        low = True
        raw = raw[1:]

    high = False
    if raw.endswith("'"):
        high = True
        raw = raw[:-1]
    elif raw.endswith("."):
        low = True
        raw = raw[:-1]

    # komal inline D(k)
    m = re.match(r"^([RGDN])\((k|K)\)$", raw)
    if m:
        raw = m.group(1).lower()

    base = raw
    if base not in _TOKEN_TO_WESTERN:
        return None

    step, alter = _TOKEN_TO_WESTERN[base]
    octave = default_octave + (1 if high else 0) - (1 if low else 0)

    acc = "♭" if alter == -1 else ("#" if alter == 1 else "")
    western_label = f"{step}{acc}{octave}"

    indian_label = tok.replace("’", "'")  # stable

    return ParsedNote(step=step, alter=alter, octave=octave, indian_label=indian_label, western_label=western_label)


def _expand_tokens_to_events(tokens: List[str]) -> List[Tuple[str, str]]:
    """
    Converts tokens into events list:
      ("note", "G")
      ("grace", "R")
      ("slur_start", "")
      ("slur_stop", "")
      ("hold", "G")   # duration doubled
    Canonical hold token: ends with ":" (e.g., "G:")
    Meend: "G~R"
    Kan: "(R)G"
    """
    events: List[Tuple[str, str]] = []
    for tok in tokens:
        t = tok.strip()
        if not t:
            continue

        # hold: canonical marker
        if t.endswith(":"):
            events.append(("hold", t[:-1]))
            continue

        # kan
        km = _RE_KAN.match(t)
        if km:
            events.append(("grace", km.group(1)))
            events.append(("note", km.group(2)))
            continue

        # meend
        mm = _RE_MEEND.match(t)
        if mm:
            events.append(("note", mm.group(1)))
            events.append(("slur_start", ""))
            events.append(("note", mm.group(2)))
            events.append(("slur_stop", ""))
            continue

        events.append(("note", t))

    return events


def make_musicxml_per_song(
    book_title: str,
    songs: List[Dict[str, Any]],
    base_dir: Path,
    outdir: Path,
    *,
    divisions: int = 2,
    beats: int = 4,
    beat_type: int = 4,
    default_octave: int = 4,
    default_note_duration: int = 1,  # in divisions units; 1 = eighth when divisions=2
) -> None:
    """
    Creates one MusicXML per song in outdir:
      <song_id>.musicxml

    Notes:
    - Rhythm: every token becomes an eighth-note by default (safe).
      Holds ("G:") double duration.
    - Meend -> slur between two notes.
    - Kan -> grace note before main note.
    - Lyrics:
      lyric 1 = Indian token
      lyric 2 = Western pitch+octave
    """
    outdir.mkdir(parents=True, exist_ok=True)

    notes_per_measure = beats * divisions  # if duration=1 per note

    def w(s: str) -> str:
        return xesc(s)

    for s in songs:
        title = s["title"]
        song_id = s["id"]
        out_path = outdir / f"{song_id}.musicxml"

        xml: List[str] = []
        xml.append('<?xml version="1.0" encoding="UTF-8" standalone="no"?>')
        xml.append('<!DOCTYPE score-partwise PUBLIC "-//Recordare//DTD MusicXML 3.1 Partwise//EN" '
                   '"http://www.musicxml.org/dtds/partwise.dtd">')
        xml.append('<score-partwise version="3.1">')
        xml.append(f'  <work><work-title>{w(title)}</work-title></work>')
        xml.append('  <identification><encoding><software>Songbook Pipeline</software></encoding></identification>')
        xml.append('  <part-list><score-part id="P1"><part-name>Lead</part-name></score-part></part-list>')
        xml.append('  <part id="P1">')

        measure_no = 1
        current_units = 0

        def start_measure(first: bool = False) -> None:
            nonlocal current_units
            xml.append(f'    <measure number="{measure_no}">')
            if first:
                xml.append('      <attributes>')
                xml.append(f'        <divisions>{divisions}</divisions>')
                xml.append('        <key><fifths>0</fifths></key>')  # SA=C
                xml.append(f'        <time><beats>{beats}</beats><beat-type>{beat_type}</beat-type></time>')
                xml.append('        <clef><sign>G</sign><line>2</line></clef>')
                xml.append('      </attributes>')
                # Place global description below the staff.
                xml.append('      <direction placement="below"><direction-type><words>'
                           'Lyric1=Indian (Sargam), Lyric2=Western (Pitch+Oct). '
                           'Ornaments: meend "~" as slur, kan "(x)y" as grace, hold ":" as longer duration.'
                           '</words></direction-type></direction>')
            current_units = 0

        def end_measure() -> None:
            xml.append('    </measure>')

        def emit_note(pn: ParsedNote, duration: int, note_type: str,
                      is_grace: bool = False, slur: Optional[str] = None) -> None:
            """
            slur: "start" | "stop" | None
            """
            xml.append('      <note>')
            if is_grace:
                xml.append('        <grace/>')
            xml.append('        <pitch>')
            xml.append(f'          <step>{pn.step}</step>')
            if pn.alter != 0:
                xml.append(f'          <alter>{pn.alter}</alter>')
            xml.append(f'          <octave>{pn.octave}</octave>')
            xml.append('        </pitch>')
            if not is_grace:
                xml.append(f'        <duration>{duration}</duration>')
                xml.append(f'        <type>{note_type}</type>')
            else:
                # Grace notes typically omit duration; MuseScore accepts this.
                xml.append('        <type>eighth</type>')
            if slur:
                xml.append('        <notations>')
                xml.append(f'          <slur type="{slur}"/>')
                xml.append('        </notations>')

            # Lyrics lines
            xml.append(f'        <lyric number="1"><text>{w(pn.indian_label)}</text></lyric>')
            xml.append(f'        <lyric number="2"><text>{w(pn.western_label)}</text></lyric>')

            xml.append('      </note>')

        def duration_to_type(dur_units: int) -> str:
            """
            Map duration units (divisions-based) to type label.
            With divisions=2:
              1 -> eighth
              2 -> quarter
              4 -> half
              8 -> whole
            """
            if dur_units <= 1:
                return "eighth"
            if dur_units == 2:
                return "quarter"
            if dur_units == 4:
                return "half"
            if dur_units == 8:
                return "whole"
            # fallback
            return "quarter"

        start_measure(first=True)

        sections = s.get("sections", []) or []
        for section in sections:
            sec_name = (section.get("name") or "").strip()
            if sec_name:
                # Place section label below the staff.
                xml.append('      <direction placement="below"><direction-type><words>'
                           f'{w(sec_name)}</words></direction-type></direction>')

            for line in (section.get("lines") or []):
                # Direction for lyrics phrase (optional)
                phrase = str(line.get("lyrics", "") or "").strip()
                if phrase:
                    # Place lyrics/phrase label below the staff (instead of above).
                    xml.append('      <direction placement="below"><direction-type><words>'
                               f'{w(phrase)}</words></direction-type></direction>')

                tokens = line.get("tokens")
                if not tokens:
                    tokens = tokenize_from_indian(str(line.get("indian", "") or ""))

                events = _expand_tokens_to_events(tokens)

                slur_pending = None  # None/"start"/"stop" applied to next note emission where relevant

                for ev, payload in events:
                    if ev == "slur_start":
                        slur_pending = "start"
                        continue
                    if ev == "slur_stop":
                        slur_pending = "stop"
                        continue

                    if ev in ("note", "hold", "grace"):
                        dur = default_note_duration
                        if ev == "hold":
                            dur = default_note_duration * 2

                        pn = _parse_note_token(payload, default_octave=default_octave)
                        if not pn:
                            continue

                        # New measure if this note would exceed
                        if current_units + dur > notes_per_measure:
                            end_measure()
                            measure_no += 1
                            start_measure(first=False)

                        note_type = duration_to_type(dur)
                        is_grace = (ev == "grace")
                        # For grace, don't advance current_units
                        emit_note(pn, dur, note_type, is_grace=is_grace, slur=slur_pending)
                        if not is_grace:
                            current_units += dur
                        slur_pending = None

        end_measure()
        xml.append('  </part>')
        xml.append('</score-partwise>')

        out_path.write_text("\n".join(xml), encoding="utf-8")


# -----------------------------
# CLI
# -----------------------------
def main() -> None:
    ap = argparse.ArgumentParser(description="Generate mobile PDF / EPUB / MusicXML songbook from songs.json")
    ap.add_argument("--input", default="songs.json", help="Path to songs.json")
    ap.add_argument("--outdir", default="output", help="Output directory")
    ap.add_argument("--format", default="both", choices=["pdf", "epub", "docx", "both", "musicxml", "all"], help="Output format")
    ap.add_argument("--page", default="A5", choices=["A5", "LETTER"], help="PDF page size")
    ap.add_argument("--bg-opacity", type=float, default=0.10, help="PDF background opacity (0.0-0.2 recommended)")
    ap.add_argument(
        "--pdf-bg-mode",
        default="cover",
        choices=["cover", "tile", "contain"],
        help="PDF background rendering: cover=fill/crop, tile=repeat, contain=fit inside",
    )
    ap.add_argument(
        "--pdf-emoji-font",
        default="",
        help="Optional path to an emoji-capable .ttf font to embed for 🎤 ✍ 🎼 (e.g., fonts/Symbola.ttf).",
    )
    ap.add_argument(
        "--pdf-variants",
        default="all",
        choices=["all", "combined", "indian", "western"],
        help="Which PDF(s) to generate: combined (lyrics+indian+western), indian (lyrics+indian), western (lyrics+western), all (all three).",
    )
    ap.add_argument(
        "--normalize-export",
        action="store_true",
        help="Write `export: true` into any songs missing the key (does not overwrite existing true/false).",
    )
    ap.add_argument("--epub-bg-opacity", type=float, default=0.10, help="EPUB background opacity (0.0-0.2 recommended)")
    ap.add_argument("--musicxml-divisions", type=int, default=2, help="MusicXML divisions")
    ap.add_argument("--musicxml-beats", type=int, default=4, help="MusicXML beats")
    ap.add_argument("--musicxml-beat-type", type=int, default=4, help="MusicXML beat-type")
    ap.add_argument("--musicxml-octave", type=int, default=4, help="Default octave for SA")
    args = ap.parse_args()

    base_dir = Path(".").resolve()
    json_path = (base_dir / args.input).resolve()
    outdir = (base_dir / args.outdir).resolve()
    outdir.mkdir(parents=True, exist_ok=True)

    if args.normalize_export:
        n = normalize_export_flags(json_path)
        print(f"Normalized export flags in {json_path} (updated {n} song(s))")

    book_title, book_meta, songs = load_songbook(json_path)

    # outputs
    if args.format in ("pdf", "both", "all"):
        base_name = book_title.replace(" ", "_")
        variants: List[Tuple[str, bool, bool, str]] = []
        if args.pdf_variants == "combined":
            variants = [("combined", True, True, f"{base_name}.pdf")]
        elif args.pdf_variants == "indian":
            variants = [("indian", True, False, f"{base_name}_Indian.pdf")]
        elif args.pdf_variants == "western":
            variants = [("western", False, True, f"{base_name}_Western.pdf")]
        else:
            variants = [
                ("combined", True, True, f"{base_name}.pdf"),
                ("indian", True, False, f"{base_name}_Indian.pdf"),
                ("western", False, True, f"{base_name}_Western.pdf"),
            ]

        try:
            for _, si, sw, fn in variants:
                pdf_path = outdir / fn
                make_pdf(
                    book_title,
                    book_meta,
                    songs,
                    base_dir,
                    pdf_path,
                    bg_opacity=args.bg_opacity,
                    page=args.page,
                    pdf_bg_mode=args.pdf_bg_mode,
                    pdf_emoji_font=args.pdf_emoji_font,
                    show_indian=si,
                    show_western=sw,
                )
                print("PDF:", pdf_path)
        except SystemExit:
            # For "--format all" we still want to produce other outputs if PDF deps are missing.
            if args.format == "all":
                print("PDF: skipped (missing reportlab)")
            else:
                raise

    if args.format in ("epub", "both", "all"):
        epub_path = outdir / f"{book_title.replace(' ', '_')}.epub"
        make_epub(book_title, book_meta, songs, base_dir, epub_path, bg_opacity=args.epub_bg_opacity)
        print("EPUB:", epub_path)

    if args.format in ("docx", "all"):
        docx_path = outdir / f"{book_title.replace(' ', '_')}.docx"
        make_docx(book_title, songs, base_dir, docx_path)
        print("DOCX:", docx_path)

    if args.format in ("musicxml", "all"):
        make_musicxml_per_song(
            book_title,
            songs,
            base_dir,
            outdir,
            divisions=args.musicxml_divisions,
            beats=args.musicxml_beats,
            beat_type=args.musicxml_beat_type,
            default_octave=args.musicxml_octave,
        )
        print("MUSICXML: generated per song in", outdir)


if __name__ == "__main__":
    main()
